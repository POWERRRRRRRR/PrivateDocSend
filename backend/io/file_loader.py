"""File I/O for txt, markdown, and docx documents."""

import os


def load_file(path: str) -> str:
    """Load a document and return its text content.

    Supports: .txt, .md, .docx
    """
    ext = os.path.splitext(path)[1].lower()

    if ext in (".txt", ".md", ".markdown"):
        with open(path, "r", encoding="utf-8") as f:
            return f.read()
    elif ext == ".docx":
        return _load_docx(path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _load_docx(path: str) -> str:
    """Extract text from a .docx file using python-docx."""
    from docx import Document

    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n\n".join(paragraphs)


def save_file(path: str, content: str) -> None:
    """Save text content to a file.

    Supports: .txt, .md, .docx
    """
    ext = os.path.splitext(path)[1].lower()

    if ext in (".txt", ".md", ".markdown"):
        with open(path, "w", encoding="utf-8") as f:
            f.write(content)
    elif ext == ".docx":
        _save_docx(path, content)
    else:
        raise ValueError(f"Unsupported file format: {ext}")


def _save_docx(path: str, content: str) -> None:
    """Write text content to a .docx file."""
    from docx import Document

    doc = Document()
    for paragraph in content.split("\n\n"):
        doc.add_paragraph(paragraph)
    doc.save(path)
